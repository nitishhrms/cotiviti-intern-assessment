"""
Generate the Cotiviti Intern Assessment written report (.docx).

Topic 1 - Clinical Natural Language Technology for Health Care:
Past, Present, & Future Approaches (NLP, OCR, Computer Vision, LLM, LMM).

Author: Nitish Kumar - San Jose State University
"""

import os

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- Brand palette (Cotiviti-inspired) -------------------------------------
PURPLE = RGBColor(0x4B, 0x2E, 0x83)      # deep purple
MAGENTA = RGBColor(0xC4, 0x00, 0x7A)     # accent magenta
DARK = RGBColor(0x22, 0x22, 0x22)        # body text
GREY = RGBColor(0x66, 0x66, 0x66)        # secondary text

FONT = "Calibri"


def set_cell_background(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_bottom_border(paragraph, color="4B2E83", size="12"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def style_body(doc):
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(10.5)
    style.font.color.rgb = DARK
    pf = style.paragraph_format
    pf.line_spacing = 1.15
    pf.space_after = Pt(6)


def heading(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    run = p.add_run(text.upper())
    run.font.name = FONT
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = PURPLE
    run.font.all_caps = True
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    add_bottom_border(p, color="C4007A", size="8")
    return p


def body(doc, text, justify=True):
    p = doc.add_paragraph(text)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def bullet(doc, bold_lead, rest):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(bold_lead)
    r.font.bold = True
    r.font.color.rgb = PURPLE
    p.add_run(rest)
    p.paragraph_format.space_after = Pt(4)
    return p


def figure(doc, filename, caption, width=6.0):
    path = os.path.join(os.path.dirname(__file__), filename)
    if not os.path.exists(path):
        return
    pic = doc.add_paragraph()
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic.paragraph_format.space_before = Pt(4)
    pic.paragraph_format.space_after = Pt(2)
    pic.add_run().add_picture(path, width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption)
    cr.font.size = Pt(9)
    cr.font.italic = True
    cr.font.color.rgb = GREY
    cap.paragraph_format.space_after = Pt(6)
    return cap


doc = Document()
style_body(doc)

# margins
for section in doc.sections:
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

# ---- Title block -----------------------------------------------------------
brand = doc.add_paragraph()
brand.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = brand.add_run("COTIVITI  ")
r.font.size = Pt(12)
r.font.bold = True
r.font.color.rgb = PURPLE
r2 = brand.add_run("| Intern Assessment")
r2.font.size = Pt(12)
r2.font.color.rgb = GREY

title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(6)
tr = title.add_run("Clinical Natural Language Technology for Health Care")
tr.font.size = Pt(20)
tr.font.bold = True
tr.font.color.rgb = PURPLE
tr.font.name = FONT

sub = doc.add_paragraph()
sr = sub.add_run("Past, Present, and Future Approaches: NLP, OCR, Computer Vision, LLMs & LMMs")
sr.font.size = Pt(12)
sr.font.italic = True
sr.font.color.rgb = MAGENTA

meta = doc.add_paragraph()
mr = meta.add_run("Nitish Kumar  •  San Jose State University  •  Topic 1")
mr.font.size = Pt(10)
mr.font.color.rgb = GREY
add_bottom_border(meta, color="4B2E83", size="12")
meta.paragraph_format.space_after = Pt(8)

# ---- Body ------------------------------------------------------------------
heading(doc, "Defining the Concept")
body(doc,
     "Clinical Natural Language Technology is a family of computational methods that convert unstructured "
     "medical information, such as physician notes, discharge summaries, scanned faxes, and diagnostic "
     "images, into clean, structured, machine-readable data. An estimated 80% of healthcare data is "
     "unstructured, locked in free text and pictures that systems cannot directly search, code, or audit. "
     "Five complementary technologies do this work. Natural Language Processing (NLP) extracts facts and "
     "clinical entities from text. Optical Character Recognition (OCR) digitizes scanned and faxed "
     "documents. Computer Vision (CV) interprets medical images. Large Language Models (LLMs) and Large "
     "Multimodal Models (LMMs) reason over text and images together. Collectively, they translate "
     "narrative medical content into the codes and supporting evidence that payment-integrity, "
     "risk-adjustment, and quality teams rely on every day.")

heading(doc, "Trends: Past, Present, and Future")
body(doc,
     "The field has evolved through three overlapping eras, summarized in Figure 1. From roughly 2000 to "
     "2017, rule-based systems such as cTAKES and MetaMap mapped text to standard terminologies including "
     "UMLS, SNOMED CT, ICD, and LOINC. These systems were precise and transparent, but brittle and "
     "expensive to maintain. From 2018 to 2023, transformer models trained on clinical text, such as "
     "BioBERT, ClinicalBERT, and GatorTron, sharply improved entity recognition, de-identification, and "
     "phenotyping, while vision models such as DenseNet reached radiologist-level performance on chest-film "
     "findings (Rajpurkar et al., 2017). Since 2024, the frontier has become generative and multimodal. "
     "Clinical LLMs such as Med-PaLM 2 encode broad medical knowledge, and LMMs interpret an image together "
     "with its clinical context, so a single model can read an X-ray and draft a report. Retrieval-Augmented "
     "Generation (RAG) has become the standard safeguard, grounding each answer in a trusted source to curb "
     "hallucination (confident but false output). In regulated healthcare, that grounding is a prerequisite, "
     "not an option.")

figure(doc, "trends_timeline.png",
       "Figure 1. Evolution of clinical natural language technology across three eras.",
       width=6.4)

heading(doc, "Opportunities")
bullet(doc, "Automated coding and payment integrity: ",
       "extract diagnoses, procedures, and supporting evidence from notes to surface miscoding, upcoding, "
       "and documentation gaps at scale. This aligns directly with Cotiviti's prepay and postpay review and "
       "risk-adjustment work.")
bullet(doc, "Multimodal claim validation: ",
       "draft radiology and pathology reports and compare imaging against notes and billed claims, "
       "highlighting findings that support or contradict a submission.")
bullet(doc, "Grounded and auditable AI: ",
       "Retrieval-Augmented Generation surfaces the exact source text behind every determination, producing "
       "the audit trail that clinical and regulatory review require.")

heading(doc, "Threats and Risks")
bullet(doc, "Hallucination and safety: ",
       "models can produce confident but false statements, which is unacceptable in payment and care "
       "decisions and must be controlled through grounding and human review.")
bullet(doc, "Privacy and compliance: ",
       "protected health information in text and images demands rigorous de-identification, private "
       "HIPAA-aligned hosting, and strict data governance.")
bullet(doc, "Bias, drift, and regulation: ",
       "performance can vary across patient subgroups and degrade as coding rules change, all under "
       "expanding FDA and payer oversight of clinical AI.")

heading(doc, "Recommended Strategic Actions for Cotiviti")
body(doc,
     "Three actions, illustrated in Figure 2, would let Cotiviti capture these opportunities while managing "
     "the risks. First, build a grounded clinical-NLP layer that pairs medical encoders such as ClinicalBERT "
     "or GatorTron with a Retrieval-Augmented Generation index of coding policies and clinical guidelines, "
     "so every code or flag the system produces links back to a named, auditable source. Second, pilot "
     "multimodal document understanding that combines OCR, NLP, and computer vision to reconcile imaging and "
     "notes against submitted claims, catching gaps that text-only review misses. Third, keep experts in the "
     "loop by versioning models, monitoring calibration and drift, and routing low-confidence cases to human "
     "reviewers, so automation augments rather than replaces them. Together, these steps give Cotiviti speed "
     "while meeting the accuracy, transparency, and compliance bar that payment-integrity work demands.")

figure(doc, "strategy_actions.png",
       "Figure 2. Three recommended strategic actions for Cotiviti.",
       width=6.4)

heading(doc, "Proof of Concept")
body(doc,
     "The prototype with this report is a Chest X-Ray Report Generation System. It shows these ideas "
     "working from start to finish. A DenseNet-121 encoder reads an X-ray in JPEG, PNG, or DICOM form. It "
     "turns the image into a feature vector. A Transformer decoder turns that vector into a draft report. "
     "The system then lays the draft out in standard radiology sections and offers it as a PDF. This is "
     "the computer vision and NLP pipeline in small form. A separate clinical agent answers questions with "
     "RAG. It turns the question into a ClinicalBERT embedding. It finds similar cases by cosine "
     "similarity. It shows each step of its reasoning, so a reviewer can check the answer. The repository "
     "also holds a larger version of that agent. That version uses ChromaDB and LangChain, with FastAPI, "
     "Celery, Prometheus, OpenTelemetry, and Docker. The public demo runs a smaller version so that it "
     "fits on free CPU hosting. The caption model is small, and it trained for only nine epochs. So the "
     "drafts it writes are short. This prototype proves the pipeline and its traceability. It does not "
     "prove diagnostic accuracy. It is a research demo. It is not FDA-cleared for clinical use.")

# ---- Page 3: References ----------------------------------------------------
doc.add_page_break()
heading(doc, "References (APA 7th Edition)")

refs = [
    "Alsentzer, E., Murphy, J., Boag, W., Weng, W.-H., Jin, D., Naumann, T., & McDermott, M. (2019). "
    "Publicly available clinical BERT embeddings. Proceedings of the 2nd Clinical Natural Language "
    "Processing Workshop, 72–78.",

    "Demner-Fushman, D., Kohli, M. D., Rosenman, M. B., Shooshan, S. E., Rodriguez, L., Antani, S., "
    "Thoma, G. R., & McDonald, C. J. (2016). Preparing a collection of radiology examinations for "
    "distribution and retrieval. Journal of the American Medical Informatics Association, 23(2), 304–310.",

    "Devlin, J., Chang, M.-W., Lee, K., & Toutanova, K. (2019). BERT: Pre-training of deep bidirectional "
    "transformers for language understanding. Proceedings of NAACL-HLT, 4171–4186.",

    "Huang, G., Liu, Z., van der Maaten, L., & Weinberger, K. Q. (2017). Densely connected convolutional "
    "networks. Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition, 4700–4708.",

    "Irvin, J., Rajpurkar, P., Ko, M., Yu, Y., Ciurea-Ilcus, S., Chute, C., ... Ng, A. Y. (2019). "
    "CheXpert: A large chest radiograph dataset with uncertainty labels and expert comparison. "
    "Proceedings of the AAAI Conference on Artificial Intelligence, 33(1), 590–597.",

    "Johnson, A. E. W., Pollard, T. J., Berkowitz, S. J., Greenbaum, N. R., Lungren, M. P., Deng, C.-Y., "
    "Mark, R. G., & Horng, S. (2019). MIMIC-CXR, a de-identified publicly available database of chest "
    "radiographs with free-text reports. Scientific Data, 6, 317.",

    "Lee, J., Yoon, W., Kim, S., Kim, D., Kim, S., So, C. H., & Kang, J. (2020). BioBERT: A pre-trained "
    "biomedical language representation model for biomedical text mining. Bioinformatics, 36(4), 1234–1240.",

    "Lewis, P., Perez, E., Piktus, A., Petroni, F., Karpukhin, V., Goyal, N., ... Kiela, D. (2020). "
    "Retrieval-augmented generation for knowledge-intensive NLP tasks. Advances in Neural Information "
    "Processing Systems, 33, 9459–9474.",

    "Rajpurkar, P., Irvin, J., Zhu, K., Yang, B., Mehta, H., Duan, T., ... Ng, A. Y. (2017). CheXNet: "
    "Radiologist-level pneumonia detection on chest X-rays with deep learning. arXiv:1711.05225.",

    "Singhal, K., Azizi, S., Tu, T., Mahdavi, S. S., Wei, J., Chung, H. W., ... Natarajan, V. (2023). "
    "Large language models encode clinical knowledge. Nature, 620, 172–180.",

    "Thirunavukarasu, A. J., Ting, D. S. J., Elangovan, K., Gutierrez, L., Tan, T. F., & Ting, D. S. W. "
    "(2023). Large language models in medicine. Nature Medicine, 29, 1930–1940.",

    "Vaswani, A., Shazeer, N., Parmar, N., Uszkoreit, J., Jones, L., Gomez, A. N., Kaiser, Ł., & "
    "Polosukhin, I. (2017). Attention is all you need. Advances in Neural Information Processing Systems, 30.",

    "Yang, X., Chen, A., PourNejatian, N., Shin, H. C., Smith, K. E., Parisien, C., ... Wu, Y. (2022). "
    "A large language model for electronic health records. npj Digital Medicine, 5, 194.",
]

for ref in refs:
    p = doc.add_paragraph(ref)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)  # hanging indent
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    for run in p.runs:
        run.font.size = Pt(10)

import os

# ---- Page 4: Appendix A — System Architecture (HLD) ------------------------
doc.add_page_break()
heading(doc, "Appendix A: System Architecture (High-Level Design)")
body(doc,
     "The proof of concept is organized in four layers. Client interfaces (a Gradio UI, Swagger, and REST "
     "clients) call a FastAPI service, which routes requests to the model, DICOM, RAG, and async services; "
     "these read from the vector store, model weights, and vocabulary.")

_hld = os.path.join(os.path.dirname(__file__), "hld_architecture.png")
if os.path.exists(_hld):
    pic_p = doc.add_paragraph()
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.paragraph_format.space_before = Pt(2)
    pic_p.paragraph_format.space_after = Pt(2)
    pic_p.add_run().add_picture(_hld, width=Inches(6.1))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run("Figure 3. High-level design of the Chest X-Ray Report Generation System.")
    cr.font.size = Pt(9)
    cr.font.italic = True
    cr.font.color.rgb = GREY

# ---- Appendix B — API Reference --------------------------------------------
heading(doc, "Appendix B: API Reference (Surface Level)")
body(doc,
     "The system exposes a REST API built with FastAPI; full interactive documentation is auto-generated "
     "at /docs (Swagger). The core endpoints are summarized below.")

api = [
    ("SYSTEM", [
        ("GET", "/health", "Liveness probe and model-loaded status"),
        ("GET", "/metadata", "Model architecture details (vocab size, dimensions, dataset)"),
        ("GET", "/metrics", "Prometheus metrics (request counts, latency histogram)"),
    ]),
    ("INFERENCE", [
        ("POST", "/predict", "Generate a report from a JPEG or PNG image"),
        ("POST", "/predict/dicom", "Generate a report from a DICOM file, with windowing"),
        ("POST", "/preview/dicom", "Return a windowed DICOM preview image (JPEG)"),
        ("GET", "/task/{task_id}", "Poll the status and result of an async job"),
    ]),
    ("AGENT", [
        ("POST", "/agent/query", "Grounded clinical question answering (optional reasoning trace)"),
        ("POST", "/agent/search", "Semantic search over indexed reports"),
        ("POST", "/agent/index", "Add reports to the vector store"),
        ("DELETE", "/agent/index", "Delete the vector store collection"),
        ("GET", "/agent/stats", "Vector store statistics (backend and count)"),
    ]),
]

table = doc.add_table(rows=0, cols=3)
table.style = "Table Grid"
table.autofit = False
col_w = [Inches(0.9), Inches(2.15), Inches(3.65)]


def _set_widths(row):
    for cell, w in zip(row.cells, col_w):
        cell.width = w


def _fmt(cell, text, *, bold=False, color=DARK, size=9.5, mono=False, align=None):
    cell.paragraphs[0].alignment = align
    run = cell.paragraphs[0].add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = "Consolas" if mono else FONT
    cell.paragraphs[0].paragraph_format.space_after = Pt(1)
    cell.paragraphs[0].paragraph_format.space_before = Pt(1)


# header row
hdr = table.add_row()
_set_widths(hdr)
for c, t in zip(hdr.cells, ("Method", "Endpoint", "Description")):
    set_cell_background(c, "4B2E83")
    _fmt(c, t, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

for group, rows in api:
    grp = table.add_row()
    _set_widths(grp)
    gcell = grp.cells[0].merge(grp.cells[1]).merge(grp.cells[2])
    set_cell_background(gcell, "EFE7FA")
    _fmt(gcell, group, bold=True, color=PURPLE, size=9)
    for method, ep, desc in rows:
        r = table.add_row()
        _set_widths(r)
        _fmt(r.cells[0], method, bold=True, color=MAGENTA, size=9, mono=True)
        _fmt(r.cells[1], ep, color=DARK, size=9, mono=True)
        _fmt(r.cells[2], desc, color=DARK, size=9)

note = doc.add_paragraph()
nr = note.add_run("Interactive requests support both synchronous and asynchronous (Celery) modes; "
                  "responses include the generated report, a confidence score, and processing time.")
nr.font.size = Pt(9)
nr.font.italic = True
nr.font.color.rgb = GREY
note.paragraph_format.space_before = Pt(6)

out = os.path.join(os.path.dirname(__file__), "Nitish_Kumar_Report.docx")
doc.save(out)
print("Saved:", out)
