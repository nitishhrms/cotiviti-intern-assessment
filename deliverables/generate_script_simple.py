"""
Generate a plain-language demo video speaking script as a formatted .docx.

This is a simpler alternative to generate_script.py: short sentences, common
words, and an easy rhythm to read out loud. It also covers the structured
report output and the model status badge added to the Space app.

Author: Nitish Kumar - San Jose State University
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

PURPLE = RGBColor(0x4B, 0x2E, 0x83)
MAGENTA = RGBColor(0xC4, 0x00, 0x7A)
DARK = RGBColor(0x22, 0x22, 0x22)
GREY = RGBColor(0x66, 0x66, 0x66)
FONT = "Calibri"


def bottom_border(p, color="C4007A", size="8"):
    pPr = p._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single"); b.set(qn("w:sz"), size)
    b.set(qn("w:space"), "4"); b.set(qn("w:color"), color)
    bdr.append(b); pPr.append(bdr)


def shade(p, fill):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill); pPr.append(shd)


doc = Document()
st = doc.styles["Normal"]
st.font.name = FONT
st.font.size = Pt(11)
st.font.color.rgb = DARK
st.paragraph_format.line_spacing = 1.15
st.paragraph_format.space_after = Pt(6)

for s in doc.sections:
    s.top_margin = Inches(0.7); s.bottom_margin = Inches(0.7)
    s.left_margin = Inches(0.9); s.right_margin = Inches(0.9)


def heading(text, color=PURPLE, size=13):
    p = doc.add_paragraph()
    r = p.add_run(text.upper())
    r.font.bold = True; r.font.size = Pt(size); r.font.color.rgb = color
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
    bottom_border(p)
    return p


def cue(text):
    p = doc.add_paragraph()
    shade(p, "EFE7FA")
    r = p.add_run(text)
    r.font.italic = True; r.font.size = Pt(10); r.font.color.rgb = PURPLE
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.05)
    return p


def line(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.font.size = Pt(11.5)
    p.paragraph_format.space_after = Pt(6)
    return p


def slidetag(tag, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(tag + "  ")
    r.font.bold = True; r.font.color.rgb = MAGENTA; r.font.size = Pt(10.5)
    r2 = p.add_run(text)
    r2.font.italic = True; r2.font.color.rgb = GREY; r2.font.size = Pt(10.5)
    return p


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    p.paragraph_format.space_after = Pt(3)
    return p


# ---- Title -----------------------------------------------------------------
t = doc.add_paragraph()
tr = t.add_run("Demo Video - Speaking Script (Simple Version)")
tr.font.bold = True; tr.font.size = Pt(20); tr.font.color.rgb = PURPLE
sub = doc.add_paragraph()
sr = sub.add_run("Cotiviti Intern Assessment  |  Nitish Kumar  |  San Jose State University")
sr.font.size = Pt(11); sr.font.color.rgb = GREY
info = doc.add_paragraph()
ir = info.add_run("Topic 1: Clinical Natural Language Technology  ·  Target length: 5 to 6 minutes  ·  "
                  "Short sentences, simple words.  Bracketed text = what you do; plain text = what you say.")
ir.font.size = Pt(10); ir.font.italic = True; ir.font.color.rgb = GREY
bottom_border(info, color="4B2E83", size="12")
info.paragraph_format.space_after = Pt(8)

# ---- Segment 1 -------------------------------------------------------------
heading("Segment 1 - Intro   (on camera, about 30 seconds)")
cue("[Look straight at the camera. Sit up. Smile. Speak slowly.]")
line("Hello. My name is Nitish Kumar. I am a student at San Jose State University.")
line("For the Cotiviti intern assessment, I chose Topic One. It is about clinical "
     "natural language technology for health care.")
line("I will show you three things. First, my report. Second, my working demo. "
     "Third, what I think Cotiviti should do next.")
line("Let us start.")
cue("[Now share your screen and open the PowerPoint. Keep your webcam on if you can.]")

# ---- Segment 2 -------------------------------------------------------------
heading("Segment 2 - Report and slides   (share slides, about 2 minutes)")

slidetag("[Slide 1 - Title]", "stay here for about 10 seconds")
line("This is my topic. It is about how computers read medical text and medical images.")

slidetag("[Slide 2 - The problem]", "")
line("Here is the main problem. About eighty percent of health data is unstructured. "
     "It sits inside notes, reports, and scans.")
line("Computers cannot use that data easily. So people read it by hand. "
     "That is slow, and it costs a lot of money.")

slidetag("[Slide 3 - Trends]", "")
line("The field grew in three steps.")
line("First came simple rules and word lists. They were exact, but they broke very easily.")
line("Then came models like BERT. They learned from data. They understood context much better.")
line("Now we have large language models. They can read text and images together.")

slidetag("[Slide 4 - Opportunities and risks]", "")
line("There are big opportunities here. Faster medical coding. Faster report writing. "
     "And AI that you can actually check.")
line("But there are real risks too. A model can make up facts. Patient data must stay private. "
     "And a model can be unfair to some groups of people.")

slidetag("[Slide 5 - Recommendations]", "")
line("So here is my advice for Cotiviti.")
line("One. Always ground the AI in a real source, so you can trace every answer.")
line("Two. Keep a human in the loop for anything clinical.")
line("Three. Watch the models over time, because data changes.")

slidetag("[Slide 6 - Proof of concept]", "")
line("To test these ideas, I built a small working demo.")
line("An X-ray image goes in. A vision model reads the image. "
     "A text model writes a draft report. Then an agent explains the result.")

slidetag("[Slide 7 - Architecture]", "point at the numbers as you talk")
line("This is the design of the system. It has four layers.")
line("Clients at the top. Then the API. Then the services. And the data at the bottom.")
line("The numbers, one to five, show how one request moves through the system.")
line("The orange dashed line is the background path. Long jobs run there, "
     "so the user does not have to wait.")

slidetag("[Slide 8 - API]", "")
line("The system has a clean REST API. It has three groups. "
     "System, inference, and agent.")
line("Now let me show you the real demo.")
cue("[Switch your screen share to the Hugging Face Space in your browser.]")

# ---- Segment 3 -------------------------------------------------------------
heading("Segment 3 - Live demo   (share the browser, about 2.5 minutes)")
line("This is my project running live on Hugging Face Spaces.")

slidetag("[Point at the green status line]", "")
line("First, look at this green line at the top. It tells me the model is loaded. "
     "It shows the size, the device, and the checkpoint.")
line("If the weights were missing, this line would turn red, and the button would be off. "
     "So you always know if the system is healthy.")

slidetag("[Tab 1 - Report Generation. Upload an X-ray]", "")
line("Now I will upload a chest X-ray. I click generate.")
line("On the right you see the finished report. It has clear sections. "
     "Indication. Technique. Findings. And Impression.")
line("I want to be honest about one thing. The model itself is small, "
     "and I trained it for only nine epochs. So it writes one short line.")
line("My code takes that one line and lays it out as a proper report.")

slidetag("[Point at the provenance box at the bottom]", "important - do not skip this")
line("This part is the most important. It is the provenance box.")
line("It says exactly which words came from the model. "
     "And which words came from a template or a knowledge base.")
line("In health care you must know where every word came from. So I show it, instead of hiding it.")

slidetag("[Tab 2 - Clinical Agent]", "")
line("Next is the clinical agent. I will ask a question. "
     "For example, what does cardiomegaly indicate?")
line("The agent looks for similar cases. Then it explains the finding. "
     "Then it gives an answer.")
line("And it shows every step it took. This is called a reasoning trace. "
     "So you can check the work.")

slidetag("[Tab 3 - Knowledge Base]", "about 15 seconds")
line("This tab searches reports by meaning, not just by keywords. "
     "I can also add new reports here.")

slidetag("[Tab 4 - About]", "about 10 seconds")
line("The about tab explains how it all works. It also has a clear disclaimer. "
     "This is a research demo. It is not approved for real clinical use.")
line("So in one demo you can see the whole pipeline. "
     "An image becomes a report. And the report becomes something you can question.")
cue("[Switch back to your camera, or go to the last slide.]")

# ---- Segment 4 -------------------------------------------------------------
heading("Segment 4 - Close   (on camera, about 30 seconds)")
line("Let me finish.")
line("Most health data is stuck inside text and images. AI can unlock it. "
     "But it must be grounded, and a human must check it.")
line("My report, my code, my slides, and this video are all in my GitHub repository.")
line("Thank you very much for your time.")
cue("[Smile. Hold still for one second. Then stop the recording.]")

# ---- Tips ------------------------------------------------------------------
doc.add_page_break()
heading("Quick tips before you record")
bullet("Read the script out loud two or three times first. It will sound much smoother.")
bullet("Speak slowly. Slower always sounds more confident than faster.")
bullet("Put a small pause after each full stop. Do not rush between sentences.")
bullet("Look at the camera lens, not at your own face on the screen.")
bullet("Sit with a light in front of you, not behind you. A window works well.")
bullet("Close every other app, so no notification pops up during the share.")
bullet("Open the Hugging Face Space before you record, so it is already awake.")
bullet("Have your X-ray image ready on the desktop, so you do not hunt for it.")
bullet("If you make a mistake, stop, pause, and say the sentence again. You can cut it later.")
bullet("Record a 20 second test first. Check that your voice and screen are both clear.")

out = os.path.join(os.path.dirname(__file__), "Nitish_Kumar_Demo_Script_Simple.docx")
doc.save(out)
print("Saved:", out)
