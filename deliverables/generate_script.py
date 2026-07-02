"""
Generate the demo video speaking script as a formatted .docx.
Author: Nitish Kumar - San Jose State University
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

PURPLE = RGBColor(0x4B, 0x2E, 0x83)
MAGENTA = RGBColor(0xC4, 0x00, 0x7A)
TEAL = RGBColor(0x00, 0x7D, 0x74)
DARK = RGBColor(0x22, 0x22, 0x22)
GREY = RGBColor(0x66, 0x66, 0x66)
FONT = "Calibri"


def bottom_border(p, color="C4007A", size="8"):
    pPr = p._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single"); b.set(qn("w:sz"), size)
    b.set(qn("w:space"), "4"); b.set(qn("w:color"), color)
    bdr.append(b); pPr.append(bdr)


def shade(p, fill):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill); pPr.append(shd)


doc = Document()
st = doc.styles["Normal"]
st.font.name = FONT
st.font.size = Pt(11)
st.font.color.rgb = DARK
st.paragraph_format.line_spacing = 1.15
st.paragraph_format.space_after = Pt(6)

for s in doc.sections:
    s.top_margin = Inches(0.7); s.bottom_margin = Inches(0.7)
    s.left_margin = Inches(0.9); s.right_margin = Inches(0.9)


def heading(text, color=PURPLE, size=13):
    p = doc.add_paragraph()
    r = p.add_run(text.upper())
    r.font.bold = True; r.font.size = Pt(size); r.font.color.rgb = color
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
    bottom_border(p)
    return p


def cue(text):
    p = doc.add_paragraph()
    shade(p, "EFE7FA")
    r = p.add_run(text)
    r.font.italic = True; r.font.size = Pt(10); r.font.color.rgb = PURPLE
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.05)
    return p


def line(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.font.size = Pt(11.5)
    p.paragraph_format.space_after = Pt(6)
    return p


def slidetag(tag, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(tag + "  ")
    r.font.bold = True; r.font.color.rgb = MAGENTA; r.font.size = Pt(10.5)
    r2 = p.add_run(text)
    r2.font.italic = True; r2.font.color.rgb = GREY; r2.font.size = Pt(10.5)
    return p


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    p.paragraph_format.space_after = Pt(3)
    return p


# ---- Title -----------------------------------------------------------------
t = doc.add_paragraph()
tr = t.add_run("Demo Video - Speaking Script")
tr.font.bold = True; tr.font.size = Pt(20); tr.font.color.rgb = PURPLE
sub = doc.add_paragraph()
sr = sub.add_run("Cotiviti Intern Assessment  |  Nitish Kumar  |  San Jose State University")
sr.font.size = Pt(11); sr.font.color.rgb = GREY
info = doc.add_paragraph()
ir = info.add_run("Topic 1: Clinical Natural Language Technology  ·  Target length: 5.5 to 6.5 minutes  ·  "
                  "Bracketed text = stage directions; plain text = what you say.")
ir.font.size = Pt(10); ir.font.italic = True; ir.font.color.rgb = GREY
bottom_border(info, color="4B2E83", size="12")
info.paragraph_format.space_after = Pt(8)

# ---- Segment 1 -------------------------------------------------------------
heading("Segment 1 - Intro   (on camera, ~30 sec)")
cue("[ON CAMERA. Look at the lens, relaxed, smile.]")
line("Hi, my name is Nitish Kumar, and I'm a student at San Jose State University.")
line("For the Cotiviti intern assessment, I chose Topic 1: Clinical Natural Language Technology for "
     "Health Care.")
line("In the next few minutes I'll cover three things. First, a quick overview of my written report. "
     "Second, my proposal for Cotiviti. And third, a live demo of my proof of concept: an automated "
     "chest X-ray report generator with a grounded clinical question-answering agent.")
line("Let's get started.")
cue("[Switch to sharing the PowerPoint, full screen. Keep your webcam in the corner if you can.]")

# ---- Segment 2 -------------------------------------------------------------
heading("Segment 2 - Report & Slides   (share slides, ~2.5 min)")
slidetag("[Slide 1 - Title]", "")
line("This is my topic: Clinical Natural Language Technology, applied to radiology.")

slidetag("[Slide 2 - 80% unstructured]", "")
line("Here's the core problem. Around 80% of health data is unstructured. The details that matter, like "
     "diagnoses, procedures, and the proof behind every claim, are locked inside free-text notes, scanned "
     "faxes, and medical images. Computers can't directly search, code, or audit that. Clinical NLP is the "
     "set of technologies that turn all of it into structured, usable data.")

slidetag("[Slide 3 - Trends]", "")
line("The field has moved through three stages. It started with hand-written rules and word lists, which "
     "were accurate but fragile. Then came models that learn from big data, like BioBERT and ClinicalBERT. "
     "And now we're in the generative and multimodal stage, where a single model can read an X-ray and "
     "draft a report. The key technique that makes this safe is RAG, which keeps the answers grounded in "
     "real sources and cuts down made-up facts.")

slidetag("[Slide 4 - Opportunities & Threats]", "")
line("The opportunities are big: automated coding, faster report drafting, and traceable AI. But there are "
     "real risks too: models can make up facts, patient data needs strict privacy handling, and the rules "
     "keep changing. Any serious solution has to manage both sides.")

slidetag("[Slide 5 - Recommendations]", "")
line("So here's what I recommend for Cotiviti. First, build a grounded clinical-NLP layer using RAG, so "
     "every AI-generated code links back to an official source. Second, combine text, images, and scanned "
     "documents to check claims. And third, keep people in the loop, so automation supports expert "
     "reviewers instead of replacing them.")

slidetag("[Slide 6 - POC Overview]", "")
line("To prove these ideas, I built a working demo. An X-ray goes into a DenseNet vision encoder, a "
     "Transformer turns those features into a report, and a RAG agent answers clinical questions on top of it.")

slidetag("[Slide 7 - Architecture / HLD]", "")
line("At a high level, the system has four layers: the client interfaces, a FastAPI service, the processing "
     "services like the model and the RAG agent, and the data stores. It's wrapped with monitoring and "
     "Docker for deployment.")

slidetag("[Slide 8 - API]", "")
line("It exposes a clean REST API, grouped into system, inference, and agent endpoints, with automatic "
     "interactive docs.")
line("Now let me show you the actual working demo.")
cue("[Switch screenshare to your Hugging Face Space in the browser.]")

# ---- Segment 3 -------------------------------------------------------------
heading("Segment 3 - Live Demo   (share Hugging Face Space, ~2.5 min)")
line("This is my project deployed live on Hugging Face Spaces. Before I dive in, one note: the goal here "
     "is to prove the concept simply, so I'll focus on the main functions.")

slidetag("[Tab 1: Report Generation - upload an X-ray]", "")
line("First, report generation. I'll upload a chest X-ray here. The DenseNet encoder reads the image, and "
     "the Transformer decoder writes a draft radiology report. And here's the generated text. This is the "
     "computer vision plus NLP pipeline working end to end.")

slidetag("[Tab 2: Clinical Agent - ask a question]", "")
line("Next is the clinical agent. I'll ask it a clinical question, something like, \"What does cardiomegaly "
     "indicate?\" Notice it doesn't just answer. It shows a step-by-step reasoning trace, and the answer is "
     "grounded in retrieved cases. That grounding is exactly what controls hallucination, which matters a "
     "lot in healthcare.")

slidetag("[Tab 3: Knowledge Base - optional, ~15 sec]", "")
line("The knowledge base tab lets me add reports and search them by meaning, not just keywords. This is "
     "the retrieval part of the RAG system.")

slidetag("[Tab 4: System & Telemetry - optional, ~10 sec]", "")
line("And finally, there's a system tab showing live health and metrics, which is the production "
     "monitoring side.")

line("So in one demo, you can see the full pipeline: an image becomes a report, and the report becomes "
     "something you can ask questions about, with the answers grounded in real data.")
cue("[Switch back to camera, or to the closing slide.]")

# ---- Segment 4 -------------------------------------------------------------
heading("Segment 4 - Close   (on camera or Slide 10, ~30 sec)")
line("To wrap up. Clinical NLP unlocks the huge amount of health data trapped in text and images. "
     "Grounded, multimodal AI with people in the loop is the safe path for payment integrity. And this "
     "proof of concept shows the full pipeline working, from the model all the way to a deployed service.")
line("The report, the code, the slides, and this video are all in my GitHub repository. Thank you very "
     "much for your time, and for the opportunity.")
cue("[Smile, hold for one second, then stop recording.]")

# ---- Tips ------------------------------------------------------------------
heading("Delivery Tips", color=TEAL)
bullet("Practice once end to end before the real take. It removes most of the filler words.")
bullet("Have the X-ray image and the question ready so there is no fumbling on screen.")
bullet("Load the Hugging Face Space before you hit record so it is already warm (Spaces can cold-start).")
bullet("If generation is slow while recording, keep narrating (\"the model is processing the image now\"). "
       "Natural pauses are fine.")
bullet("Speak a little slower than feels natural. Confidence comes across in pacing.")
bullet("Good lighting on your face, mic close, quiet room, clean desktop with notifications off.")

out = os.path.join(os.path.dirname(__file__), "Nitish_Kumar_Demo_Script.docx")
doc.save(out)
print("Saved:", out)
